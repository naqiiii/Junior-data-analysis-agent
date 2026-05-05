"""
Python Executor Tool — sandboxed execution engine for agent-generated code.

Key features
------------
- Injects the dataset as `df` (pandas DataFrame) automatically
- Captures stdout + stderr separately
- Auto-saves any matplotlib figures produced
- Returns rich ExecutionResult including output, plots, and errors
- Handles large datasets via sampling
"""

from __future__ import annotations

import io
import os
import sys
import textwrap
import traceback
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

import matplotlib
matplotlib.use("Agg")   # must be set before pyplot import
import matplotlib.pyplot as plt

# Patch show to prevent non-interactive backend crashes
def _mock_show(*args, **kwargs):
    pass
plt.show = _mock_show
matplotlib.pyplot.show = _mock_show

import numpy as np
import pandas as pd


# ---------------------------------------------------------------------------
# Result dataclass
# ---------------------------------------------------------------------------

@dataclass
class ExecutionResult:
    success: bool
    output: str = ""
    error: str = ""
    saved_plots: List[str] = field(default_factory=list)
    variables: Dict[str, Any] = field(default_factory=dict)

    @property
    def full_output(self) -> str:
        parts = []
        if self.output.strip():
            parts.append(self.output.strip())
        if self.saved_plots:
            parts.append(f"Visualizations saved: {', '.join(self.saved_plots)}")
        if self.error.strip():
            parts.append(f"ERROR:\n{self.error.strip()}")
        return "\n\n".join(parts) if parts else "Code executed with no output."


# ---------------------------------------------------------------------------
# Executor
# ---------------------------------------------------------------------------

class PythonExecutor:
    """
    Stateful Python executor.  Each instance is tied to one session.
    Shared state (variables) persists across multiple execute() calls,
    so later steps can build on earlier ones.
    """

    def __init__(
        self,
        dataset_path: str,
        output_dir: str = "outputs",
        session_id: str = "session",
        max_rows: int = 10_000,
        sample_size: int = 5_000,
    ) -> None:
        self.dataset_path = dataset_path
        self.output_dir = output_dir
        self.session_id = session_id
        self.max_rows = max_rows
        self.sample_size = sample_size
        self._plot_counter = 0
        self._saved_plots: List[str] = []

        os.makedirs(output_dir, exist_ok=True)

        # Load dataset once; share across all steps
        self._df = self._load_dataset()

        # Persistent namespace shared across execute() calls
        self._namespace: Dict[str, Any] = self._build_namespace()

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def execute(self, code: str) -> ExecutionResult:
        """Execute a snippet of Python code and return an ExecutionResult."""
        code = self._sanitize(code)
        
        # Inject warning suppression to hide noisy library warnings from the UI
        code = "import warnings\nwarnings.filterwarnings('ignore')\n" + code
        
        plots_before = set(self._saved_plots)

        old_stdout = sys.stdout
        old_stderr = sys.stderr
        sys.stdout = captured_out = io.StringIO()
        sys.stderr = captured_err = io.StringIO()

        plt.close("all")   # start fresh each step

        try:
            exec(compile(code, "<analyst_code>", "exec"), self._namespace)   # noqa: S102
            self._save_open_figures()
            success = True
            error_text = ""
        except Exception:
            success = False
            error_text = traceback.format_exc()
        finally:
            sys.stdout = old_stdout
            sys.stderr = old_stderr

        stdout_text = captured_out.getvalue()
        stderr_text = captured_err.getvalue()

        new_plots = [p for p in self._saved_plots if p not in plots_before]

        return ExecutionResult(
            success=success,
            output=stdout_text,
            error=error_text + ("\n" + stderr_text if stderr_text.strip() else ""),
            saved_plots=new_plots,
        )

    @property
    def saved_plots(self) -> List[str]:
        return list(self._saved_plots)

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    def _load_dataset(self) -> pd.DataFrame:
        try:
            if self.dataset_path.lower().endswith(".tsv"):
                df = pd.read_csv(self.dataset_path, sep='\t')
            elif self.dataset_path.lower().endswith(".docx"):
                import docx
                doc = docx.Document(self.dataset_path)
                if not doc.tables:
                    print("[Executor] WARNING: No tables found in the docx document.")
                    return pd.DataFrame()
                table = doc.tables[0]
                data = [[cell.text for cell in row.cells] for row in table.rows]
                df = pd.DataFrame(data[1:], columns=data[0])
            else:
                df = pd.read_csv(self.dataset_path)

            if len(df) > self.max_rows:
                df = df.sample(self.sample_size, random_state=42).reset_index(drop=True)
                print(f"[Executor] Large dataset detected — sampled {self.sample_size} rows.")
            return df
        except Exception as e:
            print(f"[Executor] WARNING: Could not load dataset: {e}")
            return pd.DataFrame()

    def _build_namespace(self) -> Dict[str, Any]:
        """Shared namespace injected into every exec() call."""
        return {
            # Core libraries
            "pd": pd,
            "np": np,
            "plt": plt,
            "matplotlib": matplotlib,
            # scipy / sklearn available if installed
            "scipy": _try_import("scipy"),
            "sklearn": _try_import("sklearn"),
            "seaborn": _try_import("seaborn"),
            # Pre-loaded dataset
            "df": self._df.copy() if self._df is not None else pd.DataFrame(),
            # Session info
            "__output_dir__": self.output_dir,
            "__session_id__": self.session_id,
            # Builtins
            "print": print,
            "len": len,
            "range": range,
            "enumerate": enumerate,
            "zip": zip,
            "list": list,
            "dict": dict,
            "set": set,
            "tuple": tuple,
            "str": str,
            "int": int,
            "float": float,
            "bool": bool,
            "round": round,
            "min": min,
            "max": max,
            "sum": sum,
            "abs": abs,
            "sorted": sorted,
        }

    def _save_open_figures(self) -> None:
        """Auto-save any open matplotlib figures."""
        fig_nums = plt.get_fignums()
        for num in fig_nums:
            fig = plt.figure(num)
            plot_name = f"{self.session_id}_step_plot_{self._plot_counter:03d}.png"
            plot_path = os.path.join(self.output_dir, plot_name)
            try:
                fig.savefig(plot_path, dpi=150, bbox_inches="tight")
                self._saved_plots.append(plot_path)
                self._plot_counter += 1
            except Exception as e:
                print(f"[Executor] Could not save figure: {e}")
        plt.close("all")

    @staticmethod
    def _sanitize(code: str) -> str:
        """Remove markdown fences and dedent."""
        code = code.strip()
        if code.startswith("```"):
            lines = code.split("\n")
            # Remove first fence line (```python or ```)
            lines = lines[1:]
            # Remove last fence line
            if lines and lines[-1].strip().startswith("```"):
                lines = lines[:-1]
            code = "\n".join(lines)
        return textwrap.dedent(code)


# ---------------------------------------------------------------------------
# Utility
# ---------------------------------------------------------------------------

def _try_import(module_name: str) -> Any:
    try:
        import importlib
        return importlib.import_module(module_name)
    except ImportError:
        return None
