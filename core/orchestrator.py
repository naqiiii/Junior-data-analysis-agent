"""
Orchestrator — the central control loop for the Autonomous Data Analyst Agent.

Flow
----
1.  Load dataset → extract metadata → initialise memory
2.  Planner generates step-by-step analysis plan
3.  For each step (with retry logic):
        a. Analyst generates Python code
        b. Python Executor runs the code
        c. Critic validates output and scores it
        d. If rejected and retries remain → go to (a) with critic feedback
        e. Step marked completed or failed
4.  Final insight generation (Critic agent in synthesis mode)
5.  Persist session to disk, return full result
"""

from __future__ import annotations

import logging
import os
import time
from typing import Any, Dict, Optional

import pandas as pd
from langchain_groq import ChatGroq

from agents.analyst_agent import create_analyst_agent, generate_analysis_code
from agents.critic_agent import (
    CriticResult,
    create_critic_agent,
    generate_final_insights,
    validate_step_output,
)
from agents.planner_agent import create_planner_agent, generate_analysis_plan
from core.memory import AnalysisStep, MemorySystem
from core.metadata_extractor import extract_metadata
from tools.python_executor import PythonExecutor

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Configuration defaults (overridden by .env values)
# ---------------------------------------------------------------------------

DEFAULT_MAX_RETRIES = 3
DEFAULT_SCORE_THRESHOLD = 0.65
DEFAULT_MAX_PLAN_STEPS = 8
DEFAULT_MIN_PLAN_STEPS = 4


# ---------------------------------------------------------------------------
# Orchestrator
# ---------------------------------------------------------------------------

class Orchestrator:
    """
    Manages the complete multi-agent analysis pipeline.

    Instantiate once per session; call `run()` with a dataset path and query.
    """

    def __init__(
        self,
        groq_api_key: str,
        model: str = "llama-3.3-70b-versatile",
        temperature: float = 0.2,
        max_retries: int = DEFAULT_MAX_RETRIES,
        score_threshold: float = DEFAULT_SCORE_THRESHOLD,
        max_plan_steps: int = DEFAULT_MAX_PLAN_STEPS,
        min_plan_steps: int = DEFAULT_MIN_PLAN_STEPS,
        output_dir: str = "outputs",
        log_dir: str = "logs",
    ) -> None:
        self.max_retries = max_retries
        self.score_threshold = score_threshold
        self.max_plan_steps = max_plan_steps
        self.min_plan_steps = min_plan_steps
        self.output_dir = output_dir
        self.log_dir = log_dir

        os.makedirs(output_dir, exist_ok=True)
        os.makedirs(log_dir, exist_ok=True)

        # Shared LLM instance (all agents use the same model)
        self.llm = ChatGroq(
            model=model,
            temperature=temperature,
            api_key=groq_api_key,
            max_tokens=4096,
        )

        # Instantiate agents (CrewAI wrappers around the LLM)
        self.planner = create_planner_agent(self.llm)
        self.analyst = create_analyst_agent(self.llm)
        self.critic = create_critic_agent(self.llm, score_threshold)

        logger.info(
            f"Orchestrator initialised | model={model} | "
            f"max_retries={max_retries} | threshold={score_threshold}"
        )

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def run(self, dataset_path: str, query: str) -> Dict[str, Any]:
        """
        Execute the full multi-agent analysis pipeline.

        Parameters
        ----------
        dataset_path : Absolute or relative path to a CSV file.
        query        : Natural language analysis request from the user.

        Returns
        -------
        A dict containing session results, step details, insights, and visualisation paths.
        """
        start_time = time.time()
        memory = MemorySystem()
        memory.user_query = query
        memory.dataset_path = dataset_path

        self._log_stage("PIPELINE START", f"Query: {query} | Dataset: {dataset_path}")

        # ── Step 0: Load & validate dataset ───────────────────────────────────
        df = self._load_dataset(dataset_path)
        if df is None:
            return self._error_response("Failed to load dataset.", memory)

        # ── Step 1: Extract metadata ───────────────────────────────────────────
        self._log_stage("METADATA EXTRACTION")
        metadata = extract_metadata(df, query, dataset_path)
        memory.dataset_metadata = metadata
        logger.info(
            f"Dataset: {metadata['shape']['rows']} rows × {metadata['shape']['columns']} cols | "
            f"Domain: {metadata['inferred_domain']}"
        )

        # ── Step 2: Planner creates analysis plan ──────────────────────────────
        self._log_stage("PLANNING")
        plan = generate_analysis_plan(
            llm=self.llm,
            metadata=metadata,
            query=query,
            max_steps=self.max_plan_steps,
            min_steps=self.min_plan_steps,
        )
        memory.analysis_plan = plan
        logger.info(f"Plan generated: {len(plan)} steps")
        for i, step_desc in enumerate(plan, 1):
            logger.info(f"  Step {i}: {step_desc}")

        # ── Step 3: Initialise Python executor ─────────────────────────────────
        executor = PythonExecutor(
            dataset_path=dataset_path,
            output_dir=self.output_dir,
            session_id=memory.session_id,
        )

        # ── Step 4: Iterate through plan steps ─────────────────────────────────
        for step_idx, step_description in enumerate(plan, start=1):
            self._log_stage(f"STEP {step_idx}/{len(plan)}", step_description)
            step = AnalysisStep(step_id=step_idx, description=step_description)
            memory.add_step(step)

            completed = self._execute_step_with_retry(
                step=step,
                memory=memory,
                metadata=metadata,
                executor=executor,
                query=query,
            )

            memory.update_step(
                step_idx,
                code=step.code,
                output=step.output,
                critic_score=step.critic_score,
                critic_feedback=step.critic_feedback,
                status=step.status,
                retry_count=step.retry_count,
                visualizations=step.visualizations,
                error=step.error,
            )

            status_emoji = "✅" if completed else "❌"
            logger.info(
                f"{status_emoji} Step {step_idx} {'completed' if completed else 'failed'} "
                f"| score={step.critic_score:.2f} | retries={step.retry_count}"
            )

        # ── Step 5: Generate final insights ────────────────────────────────────
        self._log_stage("INSIGHT GENERATION")
        final_context = memory.build_final_context()
        insights = generate_final_insights(self.llm, final_context, metadata)
        memory.final_insights = insights

        # ── Step 6: Persist session ────────────────────────────────────────────
        session_file = memory.save_to_file(self.output_dir)
        logger.info(f"Session saved: {session_file}")

        elapsed = round(time.time() - start_time, 2)
        self._log_stage("PIPELINE COMPLETE", f"Elapsed: {elapsed}s")

        result = memory.to_dict()
        result["session_file"] = session_file
        result["elapsed_seconds"] = elapsed
        result["executor_plots"] = executor.saved_plots
        return result

    # ------------------------------------------------------------------
    # Core execution loop with retry logic
    # ------------------------------------------------------------------

    def _execute_step_with_retry(
        self,
        step: AnalysisStep,
        memory: MemorySystem,
        metadata: Dict[str, Any],
        executor: PythonExecutor,
        query: str,
    ) -> bool:
        """
        Execute one analysis step with up to self.max_retries attempts.

        Modifies `step` in place.  Returns True if the step succeeded.
        """
        critic_feedback: Optional[str] = None

        for attempt in range(self.max_retries + 1):
            step.status = "running"
            if attempt > 0:
                step.retry_count = attempt
                logger.info(f"  ↻ Retry {attempt}/{self.max_retries} for step {step.step_id}")

            # ── 4a: Analyst generates code ─────────────────────────────────
            context = memory.build_context_for_analyst(step.description)
            code = generate_analysis_code(
                llm=self.llm,
                step_description=step.description,
                context=context,
                metadata=metadata,
                critic_feedback=critic_feedback,
                retry_count=attempt,
            )
            step.code = code
            logger.debug(f"  Code generated ({len(code)} chars)")

            # ── 4b: Execute code ───────────────────────────────────────────
            exec_result = executor.execute(code)
            step.output = exec_result.full_output
            step.visualizations = exec_result.saved_plots
            if not exec_result.success:
                step.error = exec_result.error

            logger.debug(f"  Execution {'OK' if exec_result.success else 'FAILED'}")

            # ── 4c: Critic validates ───────────────────────────────────────
            critic_context = memory.build_context_for_critic(step)
            critic_result: CriticResult = validate_step_output(
                llm=self.llm,
                context=critic_context,
                code=step.code,
                output=step.output,
                step_description=step.description,
                user_query=query,
                score_threshold=self.score_threshold,
                execution_error=step.error,
            )

            step.critic_score = critic_result.score
            step.critic_feedback = critic_result.feedback

            logger.info(
                f"  Critic score: {critic_result.score:.2f} | "
                f"Approved: {critic_result.approved}"
            )

            if critic_result.approved:
                step.status = "completed"
                return True

            # Prepare for retry
            critic_feedback = critic_result.improvements
            logger.info(f"  Critic rejected: {critic_result.feedback[:120]}")

            if attempt == self.max_retries:
                # Exhausted retries — mark failed but keep best output
                step.status = "failed"
                return False

        step.status = "failed"
        return False

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _load_dataset(path: str) -> Optional[pd.DataFrame]:
        try:
            if path.lower().endswith(".tsv"):
                df = pd.read_csv(path, sep='\t')
            elif path.lower().endswith(".docx"):
                import docx
                doc = docx.Document(path)
                if not doc.tables:
                    logger.error("No tables found in the docx document.")
                    return None
                table = doc.tables[0]
                data = [[cell.text for cell in row.cells] for row in table.rows]
                df = pd.DataFrame(data[1:], columns=data[0])
            else:
                df = pd.read_csv(path)
            logger.info(f"Loaded dataset: {df.shape[0]} rows, {df.shape[1]} cols")
            return df
        except FileNotFoundError:
            logger.error(f"Dataset not found: {path}")
            return None
        except Exception as e:
            logger.error(f"Failed to read CSV: {e}")
            return None

    @staticmethod
    def _error_response(message: str, memory: MemorySystem) -> Dict[str, Any]:
        result = memory.to_dict()
        result["error"] = message
        result["final_insights"] = message
        return result

    @staticmethod
    def _log_stage(stage: str, detail: str = "") -> None:
        sep = "─" * 60
        msg = f"\n{sep}\n  ▶  {stage}"
        if detail:
            msg += f"\n     {detail}"
        msg += f"\n{sep}"
        logger.info(msg)
        print(msg)   # also print to console for visibility
